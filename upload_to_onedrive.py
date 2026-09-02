import io
import json
import os
import re
import sys
from datetime import datetime

import requests
from docx import Document

TENANT_ID     = os.environ["AZURE_TENANT_ID"]
CLIENT_ID     = os.environ["AZURE_CLIENT_ID"]
CLIENT_SECRET = os.environ["AZURE_CLIENT_SECRET"]
ONEDRIVE_USER = os.environ["ONEDRIVE_USER_EMAIL"]


def get_token():
    resp = requests.post(
        f"https://login.microsoftonline.com/{TENANT_ID}/oauth2/v2.0/token",
        data={
            "grant_type":    "client_credentials",
            "client_id":     CLIENT_ID,
            "client_secret": CLIENT_SECRET,
            "scope":         "https://graph.microsoft.com/.default",
        },
        timeout=30,
    )
    resp.raise_for_status()
    return resp.json()["access_token"]


def upload_docx(token, onedrive_path, docx_bytes):
    resp = requests.put(
        f"https://graph.microsoft.com/v1.0/users/{ONEDRIVE_USER}"
        f"/drive/root:/{onedrive_path}:/content",
        headers={
            "Authorization": f"Bearer {token}",
            "Content-Type":  "application/vnd.openxmlformats-officedocument"
                             ".wordprocessingml.document",
        },
        data=docx_bytes,
        timeout=60,
    )
    resp.raise_for_status()
    return resp.json()


def add_inline(para, text):
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            para.add_run(part[2:-2]).bold = True
        elif part:
            para.add_run(part)


def markdown_to_docx(md):
    doc = Document()
    for line in md.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            add_inline(doc.add_paragraph(style="List Bullet"), line[2:].strip())
        elif re.match(r"^\d+\. ", line):
            add_inline(doc.add_paragraph(style="List Number"),
                       re.sub(r"^\d+\.\s*", "", line))
        elif not line.strip():
            doc.add_paragraph()
        else:
            add_inline(doc.add_paragraph(), line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def safe(text):
    return re.sub(r"[^\w\-]", "_", text).strip("_")


def main():
    data = json.loads(sys.stdin.read())

    workflow  = data["workflow_type"]
    initials  = data["patient_initials"]
    dos       = data["date_of_service"]
    content   = data["note_content"]
    session   = data["session_date"]

    dt       = datetime.strptime(session, "%m/%d/%Y")
    year     = dt.strftime("%Y")
    month    = dt.strftime("%B")
    date_dir = dt.strftime("%m-%d-%Y")

    filename      = f"{safe(workflow)}_{safe(initials)}.docx"
    onedrive_path = f"ChartGPT Notes/{year}/{month}/{date_dir}/{filename}"

    token  = get_token()
    result = upload_docx(token, onedrive_path, markdown_to_docx(content))

    print(json.dumps({
        "success": True,
        "file":    filename,
        "folder":  f"ChartGPT Notes/{year}/{month}/{date_dir}",
        "web_url": result.get("webUrl", ""),
    }))


if __name__ == "__main__":
    main()
