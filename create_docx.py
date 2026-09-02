import io
import json
import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


# ── Design helpers ─────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False,
             color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_left_bar(paragraph, hex_color="4472C4"):
    """Add a colored left border bar (blockquote style)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def set_para_spacing(paragraph, before=0, after=4):
    pPr = paragraph._p.get_or_add_pPr()
    pSp = OxmlElement("w:spacing")
    pSp.set(qn("w:before"), str(before))
    pSp.set(qn("w:after"), str(after))
    pPr.append(pSp)


def add_section_heading(doc, text, color_rgb=(30, 67, 116)):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    set_font(run, name="Calibri", size=10, bold=True, color=color_rgb)
    set_para_spacing(p, before=160, after=60)
    # Bottom border under heading
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "%02X%02X%02X" % color_rgb)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_dictation_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, name="Calibri", size=11, italic=True, color=(89, 89, 89))
    add_left_bar(p, "4472C4")
    set_para_spacing(p, before=0, after=40)
    return p


def add_inline_note(para, text, base_size=11):
    """Add inline text with **bold** marker support."""
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_font(run, name="Calibri", size=base_size, bold=True,
                     color=(0, 0, 0))
        elif part:
            run = para.add_run(part)
            set_font(run, name="Calibri", size=base_size, color=(0, 0, 0))


def add_note_content(doc, md):
    for line in md.split("\n"):
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=2)
            set_para_spacing(p, before=120, after=40)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=3)
            set_para_spacing(p, before=80, after=40)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=4)
            set_para_spacing(p, before=60, after=20)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_note(p, line[2:].strip())
            set_para_spacing(p, before=0, after=20)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_note(p, re.sub(r"^\d+\.\s*", "", line))
            set_para_spacing(p, before=0, after=20)
        elif not line.strip():
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            add_inline_note(p, line)
            set_para_spacing(p, before=0, after=40)


# ── Document builder ───────────────────────────────────────────────────────────

def build_docx(workflow, initials, dos, dictation, note_content):
    doc = Document()

    # Remove default section margins (make slightly narrower)
    section = doc.sections[0]
    section.left_margin   = Pt(72)   # 1 inch
    section.right_margin  = Pt(72)
    section.top_margin    = Pt(72)
    section.bottom_margin = Pt(72)

    # ── Title block ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    t = title_p.add_run(f"{workflow}  ·  {initials}  ·  {dos}")
    set_font(t, name="Calibri", size=14, bold=True, color=(255, 255, 255))
    # Blue shaded background for title
    pPr = title_p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1E4374")
    pPr.append(shd)
    set_para_spacing(title_p, before=80, after=80)

    # ── Physician Dictation ──────────────────────────────────────────────────
    add_section_heading(doc, "Physician Dictation", color_rgb=(30, 67, 116))

    for segment in (dictation or "").split("\n---\n"):
        segment = segment.strip()
        if segment:
            add_dictation_para(doc, segment)

    # ── Generated Note ───────────────────────────────────────────────────────
    add_section_heading(doc, "Generated Note", color_rgb=(30, 67, 116))
    add_note_content(doc, note_content)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Utilities ──────────────────────────────────────────────────────────────────

def safe(text):
    return re.sub(r"[^\w\-]", "_", text).strip("_")


def main():
    data = json.loads(sys.stdin.read())

    if "notes" in data:
        notes = data["notes"]
    else:
        notes = [data]

    uploads = []

    for note in notes:
        workflow     = note["workflow_type"]
        initials     = note["patient_initials"]
        dos          = note.get("date_of_service", "")
        dictation    = note.get("dictation", "")
        content      = note["note_content"]
        session      = note.get("session_date", data.get("session_date"))

        dt           = datetime.strptime(session, "%m/%d/%Y")
        year         = dt.strftime("%Y")
        month        = dt.strftime("%B")
        date_dir     = dt.strftime("%m-%d-%Y")

        filename     = f"{safe(workflow)}_{safe(initials)}.docx"
        onedrive_dir = f"ChartGPT Notes/{year}/{month}/{date_dir}"
        docx_path    = f"/tmp/{filename}"

        with open(docx_path, "wb") as f:
            f.write(build_docx(workflow, initials, dos, dictation, content))

        uploads.append(f"{docx_path}|{onedrive_dir}")
        print(f"Created: {docx_path} → {onedrive_dir}/{filename}")

    with open("/tmp/uploads.txt", "w") as f:
        f.write("\n".join(uploads))


if __name__ == "__main__":
    main()
